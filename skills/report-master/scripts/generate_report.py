#!/usr/bin/env python3
"""
报告老油条 - DOCX 生成器示例
用法: python generate_report.py --title "周报" --output report.docx
"""

import argparse
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def setup_chinese_font(run, font_name='微软雅黑', size=12):
    """设置中文字体"""
    run.font.name = font_name
    run.font.size = Pt(size)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def add_title(doc, text):
    """添加标题"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    setup_chinese_font(run, '黑体', 16)
    run.bold = True
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)


def add_section(doc, title, content):
    """添加章节"""
    # 章节标题
    p = doc.add_paragraph()
    run = p.add_run(title)
    setup_chinese_font(run, '黑体', 14)
    run.bold = True
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)

    # 章节内容
    p = doc.add_paragraph()
    run = p.add_run(content)
    setup_chinese_font(run, '宋体', 12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5


def create_weekly_report(output_path, title="工作周报", week_range=""):
    """生成标准周报"""
    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    # 标题
    add_title(doc, title)

    # 日期行
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"报告周期：{week_range or '2026年第XX周'}")
    setup_chinese_font(run, '宋体', 10)

    p = doc.add_paragraph()
    run = p.add_run("报告人：______")
    setup_chinese_font(run, '宋体', 10)

    # 本周工作概述
    add_section(doc, "一、本周工作概述",
                "在本周期内，团队围绕核心目标，稳步推进各项工作。"
                "重点任务按计划有序开展，整体进度符合预期。"
                "针对工作中发现的问题，已及时组织专项讨论并制定相应改进措施。")

    # 具体工作内容
    add_section(doc, "二、具体工作完成情况",
                "1. 日常运营维护工作正常开展，系统运行平稳，未发生重大异常事件。\n"
                "2. 推进 XX 项目阶段性工作，目前已完成前期调研与方案设计，进入实施准备阶段。\n"
                "3. 完成内部流程优化工作，细化了部分操作规范，提升了工作效率。\n"
                "4. 配合相关部门完成协同任务，确保信息传递及时、准确。")

    # 存在问题与改进
    add_section(doc, "三、存在问题与改进方向",
                "1. 部分跨部门协作事项需进一步明确责任分工，待后续细化对接。\n"
                "2. 个别任务的交付节点受外部因素影响有所延迟，已制定追赶计划。")

    # 下周工作计划
    add_section(doc, "四、下周工作计划",
                "1. 继续推进 XX 项目后续工作，确保按计划完成阶段性目标。\n"
                "2. 针对本周识别的改进点，落实具体措施并跟踪执行效果。\n"
                "3. 保持日常工作节奏，加强内部沟通与协同。\n"
                "4. 根据实际需要，适时调整工作优先级。")

    # 结语
    add_section(doc, "五、总结",
                "总体而言，本周工作按计划稳步推进，各项任务有序推进。"
                "下阶段将继续保持当前工作节奏，重点关注关键节点的交付质量，"
                "确保各项工作目标顺利达成。")

    doc.save(output_path)
    print(f"✅ 报告已生成：{output_path}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="报告老油条生成器")
    parser.add_argument("--title", default="工作周报", help="文档标题")
    parser.add_argument("--week", default="", help="报告周期")
    parser.add_argument("--output", default="report.docx", help="输出文件名")
    args = parser.parse_args()

    create_weekly_report(args.output, args.title, args.week)
